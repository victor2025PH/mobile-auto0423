# -*- coding: utf-8 -*-
"""根据项目网络相关讨论整理，生成 Word 文档（网络架构与配置说明）。"""
import sys
from pathlib import Path

try:
    from docx import Document
    from docx.shared import Pt, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
except ImportError:
    print("请先执行: pip install python-docx", file=sys.stderr)
    sys.exit(1)

ROOT = Path(__file__).resolve().parent.parent
OUT = ROOT.parent / "docs" / "抖音" / "网络架构与配置说明_OpenClaw.docx"


def set_cell_shading(cell, fill_hex):
    from docx.oxml import OxmlElement
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill_hex)
    cell._tc.get_or_add_tcPr().append(shading)


def add_heading(doc, text, level=1):
    return doc.add_heading(text, level=level)


def add_para(doc, text, bold=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(11)
    return p


def main():
    doc = Document()
    sect = doc.sections[0]
    sect.page_height = Cm(29.7)
    sect.page_width = Cm(21.0)
    sect.left_margin = Cm(2.5)
    sect.right_margin = Cm(2.5)

    title = doc.add_heading("OpenClaw 网络架构与配置说明", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph(
        "本文档根据 OpenClaw 手机集群项目中的网络方案讨论整理，涵盖：主控/Worker 与「软路由」的关系、"
        "GL.iNet 硬件方案、代理池与路由器管理、VPN 与出口 IP、规模化（多机位）拓扑与设备选型、"
        "以及自建 VPS 与商用代理的配合方式。适用于运维与部署人员查阅。"
    )

    # --- 一、术语与概念 ---
    add_heading(doc, "一、术语与概念", 1)
    add_heading(doc, "1.1 什么是 GL.iNet", 2)
    doc.add_paragraph(
        "GL.iNet 是小型路由器/旅行路由器品牌，产品多采用 OpenWrt 或兼容系统，可安装 OpenClash 等插件，"
        "通过 Web 管理；部分型号提供 HTTP API（如 /rpc），便于主控程序远程下发配置。"
        "在本项目中，「软路由」在实现层面主要指此类独立硬件路由器上运行的透明代理（Clash/OpenClash），"
        "而非在 Windows 电脑上安装的软路由软件。"
    )
    add_heading(doc, "1.2 什么是 VPS", 2)
    doc.add_paragraph(
        "VPS（Virtual Private Server，虚拟专用服务器）是云厂商在物理机上划分出的虚拟机实例，"
        "具备独立公网 IP、操作系统（多为 Linux）与可远程 SSH 管理。"
        "若希望 TikTok 等应用看到的出口公网 IP 为「自有服务器 IP」，通常需要在 VPS 上搭建 VPN/代理并使流量从 VPS 网卡直连出网；"
        "若仅将 VPS 作为跳板再转发到第三方 SOCKS 等上游，则对外出口 IP 往往仍是上游服务商的 IP。"
    )
    add_heading(doc, "1.3 主控（Coordinator）与 Worker", 2)
    doc.add_paragraph(
        "主控运行 OpenClaw API（如端口 8000），可管理本机 ADB 设备并通过中间件将部分请求转发到 Worker。"
        "Worker 为插有手机的从节点。VPN 配置池支持主控向 Worker 同步（如 /vpn/pool/sync），"
        "但 GL.iNet 路由器列表（routers.json）为各机本地维护，不会自动从主控同步到 Worker；"
        "在 Worker 网段能访问到的路由器须在对应节点配置。"
    )

    # --- 二、项目内已实现的网络能力 ---
    add_heading(doc, "二、项目内已实现的网络能力", 1)
    add_heading(doc, "2.1 GL.iNet + OpenClash（代理中心）", 2)
    doc.add_paragraph(
        "架构：手机连接 GL.iNet 的 WiFi → 路由器上 OpenClash（Clash TUN）→ 上游代理 → 目标站点。"
        "主控侧提供 Router Manager：注册路由器、生成 Clash 配置、推送到路由器、代理与设备绑定、出口 IP 检测、备份回滚、轮换等。"
        "前端入口：Dashboard →「代理中心」（非「VPN 管理」）。「VPN 管理」侧重手机 v2rayNG 配置池。"
    )
    add_heading(doc, "2.2 未实现：电脑本机当网关 + 无线路由分 IP", 2)
    doc.add_paragraph(
        "当前代码不包含：在 Windows 上运行 Clash 作为整网网关、再通过普通 AP 给手机 DHCP 与透明代理的完整自动化与监控。"
        "若采用该拓扑，需自行维护路由/NAT/防火墙与监控，且与现有「代理中心」模块不对接。"
    )
    add_heading(doc, "2.3 手机侧 v2rayNG（VPN 配置池）", 2)
    doc.add_paragraph(
        "通过配置池添加 SOCKS5/HTTP 等，按设备分配（assignments），并可调用接口向手机导入配置（如 /vpn/pool/apply）。"
        "适合快速验证、单机调试或未布好 WiFi 路由时；长期多机运维通常更倾向 GL.iNet 统一出口。"
    )

    # --- 三、方案对比（如何选择）---
    add_heading(doc, "三、方案对比与选择建议", 1)
    t = doc.add_table(rows=4, cols=3)
    t.style = "Table Grid"
    hdr = t.rows[0].cells
    hdr[0].text = "方案"
    hdr[1].text = "优点"
    hdr[2].text = "适用场景"
    for c in hdr:
        set_cell_shading(c, "D9E2F3")
    rows = [
        ("GL.iNet + OpenClash", "手机连 WiFi 即可；与代理中心深度集成；主控重启不影響已连 WiFi 的上网。", "机位固定、多机、少动手机系统。"),
        ("每机 v2rayNG + 配置池", "不依赖路由；可一人一出口。", "过渡、抽检、小批量或必须按机换 IP。"),
        ("电脑当软路由 + AP", "不增加 GL 硬件时需自组。", "项目未集成；复杂度高，适合有专职网工。"),
    ]
    for i, row in enumerate(rows, start=1):
        for j, val in enumerate(row):
            t.rows[i].cells[j].text = val
    doc.add_paragraph(
        "结论：默认生产环境优先 GL.iNet（或同类 OpenWrt）+ 代理中心；v2rayNG 作补充；"
        "电脑网关方案除非有明确运维能力，否则一般不优先。"
    )

    # --- 四、主控侧配置要点（文件与路径）---
    add_heading(doc, "四、主控侧配置要点（文件与路径）", 1)
    add_heading(doc, "4.1 配置文件位置", 2)
    doc.add_paragraph(
        "项目根目录以 mobile-auto-project 为例（请按实际部署路径调整）："
    )
    items = [
        "config/cluster.yaml：角色 coordinator/worker、端口、集群密钥等。",
        "config/routers.json：GL.iNet 路由器登记（router_id、ip、port、password、proxy_ids、device_ids、国家等）。",
        "config/vpn_pool.json：VPN/代理配置池与设备分配（RouterManager 读取此路径下的池以生成路由器 Clash）。",
        "注意：部分环境下 API 写入的 vpn_pool 可能出现在 src/config/vpn_pool.json，与 config/vpn_pool.json 不一致时，应以 RouterManager 实际读取的 config/vpn_pool.json 为准，并建议两边对齐。",
    ]
    for it in items:
        doc.add_paragraph(it, style="List Bullet")

    add_heading(doc, "4.2 双机（主控两台手机）+ GL.iNet 登记示例逻辑", 2)
    doc.add_paragraph(
        "曾在配置中采用单台路由器 router-coordinator-eu，绑定两条上游代理 ID（如 proxy_1776301002612、proxy_1776301002726），"
        "设备 ID 示例：89NZVGKFD6BYUO5P（主控-01）、QSVSMRXOXWCYFIX4（主控-02）。"
        "Clash 对多上游可采用 load-balance + sticky-sessions，利于不同终端会话稳定落到不同上游。"
        "部署前须在 routers.json 中填写路由器真实 LAN IP（常见 192.168.8.1 需按实机修改）与 root 管理密码；"
        "路由器须安装 OpenClash；主控与路由器网络互通后，在代理中心执行部署或调用 POST /routers/{id}/deploy。"
    )

    # --- 五、商用代理（ip2up）与出口 IP ---
    add_heading(doc, "五、商用代理与出口 IP 说明", 1)
    doc.add_paragraph(
        "若使用服务商提供的 SOCKS5（如 eu-relay.ip2up.com 及账号密码），常见用法包括："
    )
    for it in [
        "直接写入 VPN 配置池，分配给指定 device_id，并通过 v2rayNG 导入；",
        "或写入代理池并分配给 GL.iNet，由 Clash 作为上游；",
        "若自建 VPS，在 VPS 上运行 sing-box/Xray 等，inbound 接手机，outbound 指向上游 SOCKS，则对外 IP 多为上游出口，而非 VPS 公网 IP。",
    ]:
        doc.add_paragraph(it, style="List Bullet")
    doc.add_paragraph(
        "若业务要求「应用侧看到的 IP 必须是自有 VPS」，需让流量从 VPS 网卡直连出网（如 WireGuard/sing-box Freedom 等架构），"
        "与「仅经 VPS 转发到第三方代理」是不同目标，需分开设计。"
    )

    # --- 六、规模化（约 120 台手机）---
    add_heading(doc, "六、规模化场景（约 120 台手机）", 1)
    doc.add_paragraph(
        "单机小型旅行路由无法稳定承载上百台活跃短视频终端。建议："
    )
    for it in [
        "按簇拆分：每簇约 10～20 台手机对应独立网关或独立出口，避免单 IP 堆叠与射频拥塞；",
        "无线侧使用多台专业 AP（如 UniFi、TP-Link Omada 等）分区覆盖，网关负责 NAT/代理，AP 负责桥接与射频；",
        "单层、不做墙内预埋时，仍可用短跳线：网关 → PoE 交换机 → 各 AP；",
        "「完全无线回程 Mesh」对大流量多终端场景通常弱于有线回程，需谨慎评估。",
    ]:
        doc.add_paragraph(it, style="List Bullet")

    add_heading(doc, "6.1 参考硬件组合（性能与稳定优先）", 2)
    doc.add_paragraph(
        "网关（单组或总出口之一）：GL.iNet GL-MT6000（Flint 2）—— CPU/内存较强，多 LAN，适合接交换机后再挂多台 AP。"
    )
    doc.add_paragraph(
        "交换机：24 口管理型 PoE+（如 Ubiquiti UniFi USW-24-PoE 或 TP-Link JetStream 同规格），"
        "为 8～12 台 AP 供电与汇聚预留端口。"
    )
    doc.add_paragraph(
        "AP：UniFi U6 系列或 Omada EAP650/670 等，按面积与终端密度布点；具体数量需现场勘测。"
    )

    # --- 七、菲律宾采购提示 ---
    add_heading(doc, "七、采购与渠道（菲律宾 Lazada 等）", 1)
    doc.add_paragraph(
        "GL.iNet 常见可搜型号：GL-MT3000（Beryl AX）、GL-B2200（Velica）、GL-BE3600（Slate 7）等，以实际库存为准。"
        "小体积 GL-MT300N-V2（Mango）仅适合极少机位。批量与大额订单建议核对卖家资质与保修。"
    )

    # --- 八、操作清单（部署 GL.iNet）---
    add_heading(doc, "八、GL.iNet 部署检查清单", 1)
    steps = [
        "路由器刷好/装好 OpenWrt 与 OpenClash，与主控网络互通。",
        "在 config/routers.json 填写正确 ip、port、password、proxy_ids、device_ids。",
        "config/vpn_pool.json 中含对应 proxy id 的完整代理条目。",
        "重启 OpenClaw 或重新加载路由管理器后，在代理中心执行部署。",
        "手机连接该路由器 WiFi，验证出口 IP 与业务可用性。",
        "若启用代理健康监控，确认手机公网 IP 与路由器记录一致（防泄漏）。",
    ]
    for i, s in enumerate(steps, 1):
        doc.add_paragraph(f"{i}. {s}")

    # --- 九、Dashboard ---
    add_heading(doc, "九、Web 控制台（Dashboard）操作指引", 1)
    add_heading(doc, "9.1 入口", 2)
    doc.add_paragraph(
        "浏览器访问主控（如 http://主控IP:8000/dashboard），按部署配置 API Key 或会话。"
        "左侧「核心」：「代理中心」= GL.iNet + 代理池 + 部署与映射；「VPN 管理」= 手机 v2rayNG 与配置池，二者勿混淆。"
    )
    add_heading(doc, "9.2 代理中心常见操作", 2)
    for it in [
        "添加路由器：GL.iNet IP、端口（常为 80）、root 密码、国家/城市。",
        "添加/导入代理：写入与 vpn_pool 一致的 SOCKS5/HTTP 等。",
        "分配代理与分配手机到路由器；一键部署或单台部署；预览 Clash、备份与回滚。",
        "检测出口 IP、健康监控、按路由器国家做批量地理配置。",
    ]:
        doc.add_paragraph(it, style="List Bullet")

    # --- 十、VPN 池与 v2rayNG ---
    add_heading(doc, "十、VPN 配置池与 v2rayNG 流程", 1)
    for it in [
        "向池中添加配置：POST /vpn/pool/add-proxy 等。",
        "绑定设备：POST /vpn/pool/assign（config_id + device_ids）。",
        "下发手机：POST /vpn/pool/apply（依赖 ADB 与 v2rayNG）。",
        "注意 config/vpn_pool.json 与 src/config/vpn_pool.json 宜保持一致，供 RouterManager 与 API 共用。",
    ]:
        doc.add_paragraph(it, style="List Bullet")

    add_heading(doc, "十一、Worker 节点", 1)
    doc.add_paragraph(
        "routers.json 不随主控自动同步；若路由器仅在 Worker 网段可达，须在该节点登记。"
        "vpn_pool 可由主控向 Worker 同步（视部署启用 /vpn/pool/sync）。"
    )

    add_heading(doc, "十二、故障排查摘要", 1)
    for it in [
        "部署失败：检查主控访问路由器 /rpc、密码、OpenClash 是否安装。",
        "手机 IP 与路由出口不一致：排查泄漏、DNS、误用蜂窝网络。",
        "v2rayNG 导入失败：USB 授权、超时、屏幕解锁。",
        "Worker API 404：节点代码版本需与主控对齐（OTA）。",
    ]:
        doc.add_paragraph(it, style="List Bullet")

    add_heading(doc, "十三、自建 VPS 与出口 IP", 1)
    doc.add_paragraph(
        "经 VPS 转发到上游 SOCKS：对外 IP 多为上游出口。若要求应用侧必须为 VPS 公网 IP，"
        "需 VPS 作为 NAT 直连出口或明确分流，不能默认等同「再接一层代理」。"
    )

    add_heading(doc, "十四、附录 API", 1)
    doc.add_paragraph("router_mgmt：/routers、/routers/{id}/deploy、assign-proxy、assign-device、clash-config、backups、restore 等。")
    doc.add_paragraph("vpn：/vpn/pool、add-proxy、assign、apply、sync 等。")
    doc.add_paragraph("cluster：OTA 与心跳；不替代 routers.json。")

    doc.add_paragraph("")
    p = doc.add_paragraph()
    p.add_run("文档生成说明：").bold = True
    p.add_run(
        "由 scripts/generate_network_docx.py 根据聊天记录整理；敏感信息以现场配置为准，勿提交公开仓库。"
    )
    p = doc.add_paragraph()
    p.add_run("合规提示：").bold = True
    p.add_run("请遵守当地法律与平台规则；本文为技术说明。")

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(OUT))
    print("已写入:", OUT)


if __name__ == "__main__":
    main()
